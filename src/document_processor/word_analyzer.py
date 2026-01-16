"""
Word Document Analyzer - Extract text, tables, images and analyze charts
"""

import os
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import base64


@dataclass
class ImageElement:
    """Image element extracted from document"""

    image_id: str
    image_path: str
    caption: str = ""
    description: str = ""  # AI-generated description (to be reviewed)
    page_number: Optional[int] = None
    inline_position: Optional[int] = None


@dataclass
class TableElement:
    """Table element extracted from document"""

    table_id: str
    data: List[List[str]]  # 2D array of cell text
    caption: str = ""
    description: str = ""  # AI-generated description (to be reviewed)
    row_count: int = 0
    column_count: int = 0
    page_number: Optional[int] = None


@dataclass
class TextElement:
    """Text element from document"""

    text: str
    element_type: str = "Paragraph"  # Title, Paragraph, etc.
    page_number: Optional[int] = None


@dataclass
class DocumentAnalysisResult:
    """Complete document analysis result"""

    text_elements: List[TextElement] = field(default_factory=list)
    table_elements: List[TableElement] = field(default_factory=list)
    image_elements: List[ImageElement] = field(default_factory=list)
    full_text: str = ""


class WordDocumentAnalyzer:
    """Analyze Word documents - extract text, tables, images, and analyze charts"""

    def __init__(self, api_client=None):
        """
        Initialize analyzer

        Args:
            api_client: Optional API client for AI-based image analysis
        """
        self.api_client = api_client

    def analyze_document(
        self, file_path: str, analyze_images: bool = True
    ) -> DocumentAnalysisResult:
        """
        Complete document analysis

        Args:
            file_path: Path to Word document
            analyze_images: Whether to use AI to analyze images (requires API)

        Returns:
            DocumentAnalysisResult with all extracted elements
        """
        result = DocumentAnalysisResult()

        # Extract text, tables, and images
        self._extract_elements(file_path, result)

        # Build full text
        result.full_text = self._build_full_text(result)

        # Analyze images with AI if requested
        if analyze_images and self.api_client and result.image_elements:
            self._analyze_images_with_ai(result)

        return result

    def _extract_elements(self, file_path: str, result: DocumentAnalysisResult):
        """Extract all elements from Word document"""
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import nsmap
        from docx.oxml import parse_xml
        from docx.shape import ShapeTypes

        doc = Document(file_path)

        # Create temp directory for images
        temp_dir = tempfile.mkdtemp()
        image_counter = 0
        table_counter = 0

        # 1. Extract paragraphs and inline elements
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Determine element type
                if para.style.name.startswith("Heading"):
                    elem_type = "Title"
                else:
                    elem_type = "Paragraph"

                result.text_elements.append(
                    TextElement(text=text, element_type=elem_type)
                )

        # 2. Extract tables (preserving order with text)
        for table in doc.tables:
            table_counter += 1
            table_data = []
            max_cols = 0

            # Extract table content
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
                max_cols = max(max_cols, len(row_data))

            # Create table element
            table_elem = TableElement(
                table_id=f"table_{table_counter}",
                data=table_data,
                row_count=len(table_data),
                column_count=max_cols,
                caption=self._find_table_caption(table, doc),
            )

            # Convert table to text for full_text
            table_text = self._format_table_text(table_data)
            result.text_elements.append(
                TextElement(text=table_text, element_type="Table")
            )

            result.table_elements.append(table_elem)

        # 3. Extract images
        # Get the document's relationships to find images
        rels = doc.part.rels

        for rel_id, rel in rels.items():
            if rel.reltype == RT.IMAGE:
                image_counter += 1
                image_part = doc.part.related_parts[rel_id]
                image_data = image_part.blob

                # Save image to temp file
                ext = self._get_image_extension(image_part.content_type)
                image_filename = f"image_{image_counter}{ext}"
                image_path = os.path.join(temp_dir, image_filename)

                with open(image_path, "wb") as f:
                    f.write(image_data)

                # Find caption
                caption = self._find_image_caption(rel_id, doc)

                image_elem = ImageElement(
                    image_id=f"image_{image_counter}",
                    image_path=image_path,
                    caption=caption,
                )

                result.image_elements.append(image_elem)

    def _format_table_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable text"""
        if not table_data:
            return ""

        lines = []
        for row in table_data:
            row_text = " | ".join(cell for cell in row if cell)
            if row_text:
                lines.append(row_text)
        return "\n".join(lines)

    def _find_table_caption(self, table, doc) -> str:
        """Find caption for a table"""
        # Try to find preceding paragraph that looks like a caption
        tables = doc.tables
        table_idx = -1

        for i, t in enumerate(tables):
            if t == table:
                table_idx = i
                break

        if table_idx > 0:
            prev_para = doc.paragraphs[table_idx - 1] if table_idx > 0 else None
            if prev_para:
                text = prev_para.text.strip()
                if "表" in text or "table" in text.lower():
                    return text

        return ""

    def _find_image_caption(self, rel_id: str, doc) -> str:
        """Find caption for an image"""
        # This is a simplified version - in practice, captions are often
        # in paragraphs near the image
        return ""

    def _get_image_extension(self, content_type: str) -> str:
        """Get file extension from content type"""
        extension_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/webp": ".webp",
        }
        return extension_map.get(content_type, ".png")

    def _build_full_text(self, result: DocumentAnalysisResult) -> str:
        """Build full text from all elements"""
        parts = []

        for elem in result.text_elements:
            if elem.element_type == "Table":
                parts.append(f"[表格数据]\n{elem.text}")
            else:
                parts.append(elem.text)

        return "\n\n".join(parts)

    def _analyze_images_with_ai(self, result: DocumentAnalysisResult):
        """Use AI to analyze images and generate descriptions"""
        for image_elem in result.image_elements:
            if not os.path.exists(image_elem.image_path):
                continue

            try:
                # Read image and encode to base64
                with open(image_elem.image_path, "rb") as f:
                    image_data = base64.b64encode(f.read()).decode()

                # Create analysis prompt
                prompt = """分析这张图表图片，请提供：
1. 图表类型（柱状图、折线图、散点图、饼图、热力图等）
2. 图表标题和坐标轴标签
3. 主要数据趋势和发现
4. 关键数值或统计结果
5. 图表的完整描述（能够还原图表内容）

请用简洁的中文描述。"""

                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}"
                                },
                            },
                        ],
                    }
                ]

                # Call API (adjust based on your API client)
                response = self.api_client.call_model(
                    model="gpt-4o", messages=messages, max_tokens=1000
                )

                image_elem.description = response

            except Exception as e:
                image_elem.description = f"图片分析失败: {str(e)}"

    def analyze_table_with_ai(self, table_elem: TableElement) -> str:
        """Use AI to analyze table and generate description"""
        if not self.api_client:
            return "需要API配置才能分析表格"

        try:
            # Format table data for prompt
            table_text = self._format_table_text(table_elem.data)

            prompt = f"""分析以下表格数据，请提供：
1. 表格概述（行数、列数、主要内容）
2. 主要发现和数据趋势
3. 关键统计结果或显著性
4. 表格的完整描述

表格数据：
{table_text}"""

            messages = [{"role": "user", "content": prompt}]

            response = self.api_client.call_model(
                model="gpt-4o", messages=messages, max_tokens=800
            )

            return response

        except Exception as e:
            return f"表格分析失败: {str(e)}"

    def _format_table_text(self, table_data: List[List[str]]) -> str:
        """Format table data for AI analysis"""
        if not table_data:
            return "空表格"

        # Header
        header = " | ".join(table_data[0]) if table_data else ""
        separator = "---" * len(table_data[0]) if table_data else ""
        rows = "\n".join(" | ".join(row) for row in table_data[1:] if row)

        return f"{header}\n{separator}\n{rows}"


class CorrectedContent:
    """Store corrected content from user review"""

    def __init__(self):
        self.text_content: str = ""
        self.table_descriptions: Dict[
            str, str
        ] = {}  # table_id -> corrected description
        self.image_descriptions: Dict[
            str, str
        ] = {}  # image_id -> corrected description
        self.results_data: str = ""  # Combined corrected results

    def to_context(self) -> Dict[str, Any]:
        """Convert to context dictionary for AI"""
        return {
            "text_content": self.text_content,
            "table_descriptions": self.table_descriptions,
            "image_descriptions": self.image_descriptions,
            "results_data": self.results_data,
        }


def analyze_document_with_review(
    file_path: str,
    api_client=None,
    analyze_images: bool = True,
) -> Tuple[DocumentAnalysisResult, CorrectedContent]:
    """
    Analyze document and prepare for user review

    Returns:
        Tuple of (analysis_result, corrected_content_placeholder)
    """
    analyzer = WordDocumentAnalyzer(api_client)
    result = analyzer.analyze_document(file_path, analyze_images=analyze_images)
    corrected = CorrectedContent()

    return result, corrected
